import re
from io import BytesIO
from docx import Document
from fpdf import FPDF
import fitz  # PyMuPDF
import os
import openai
import streamlit as st
from concurrent.futures import ThreadPoolExecutor
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
# Load API key from environment variable      

FONT_PATH = os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans.ttf")

# --- App Design Enhancements ---

# --- Function to display the pricing page ---
def show_pricing_page():
    st.title("Options d'abonnement")

    st.markdown(
        """
        ## Choisissez l'abonnement qui vous convient :

        ### Abonnement hebdomadaire : 15$

        * **6 résumés autorisés par semaine.**
        * Idéal pour les étudiants qui ont besoin de résumer quelques cours chaque semaine.
        * Accès à toutes les fonctionnalités de base.

        ### Abonnement mensuel : 50$

        * **20 résumés autorisés par mois.**
        * Parfait pour les étudiants ou professionnels qui ont besoin de résumer un grand nombre de documents.
        * Accès à toutes les fonctionnalités, y compris les fonctionnalités premium (à venir).

        ---

        **Avantages des abonnements :**

        * **Gain de temps considérable :** Obtenez des résumés concis et précis en quelques minutes.
        * **Meilleure compréhension :** Concentrez-vous sur les points clés de vos documents.
        * **Amélioration de la productivité :** Traitez plus d'informations en moins de temps.

        **Prêt à vous abonner ?**

        Cliquez sur le bouton ci-dessous pour choisir votre abonnement et commencer à profiter de tous les avantages !

        """
    )

# --- Sidebar ---
st.sidebar.title("Options de paiement")
st.sidebar.write("Prix par résumé : 5$")  # Display price per summary

if st.sidebar.button("Obtenir un abonnement"):
    show_pricing_page()  # Show the pricing page

# --- Main Content ---
st.title("Choix du LLM")

# --- Get OpenAI API key from user input ---
openai_api_key = st.text_input("Entrez votre clé API OpenAI :", type="password")

if openai_api_key:
    openai.api_key = openai_api_key
else:
    st.warning("Veuillez entrer votre clé API OpenAI pour continuer.")
    st.stop()  # Stop execution if no API key is provided


if openai_api_key is None:
    st.error("OpenAI API key not found. Please set it as an environment variable.")

openai.api_key = openai_api_key
# Constants# Constants
MAX_TOKENS_PER_REQUEST = 4096  # A safer limit to avoid hitting the token cap
PDF_CHUNK_SIZE = 3000  # Smaller chunk size for faster PDF processing
DOCX_CHUNK_SIZE = 6000  # Larger chunk size for DOCX files
MAX_WORKERS = 4  # Number of threads for parallel processing
MAX_CHUNK_TOKENS = 4000  # Maximum token limit per chunk (for safety)
# Function to read a .docx file and return a list of paragraphs
def read_docx(file):
    doc = Document(file)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    return paragraphs

# Optimized function to read a PDF file and return a list of paragraphs using fitz (PyMuPDF)
def read_pdf(file):
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    
    # Iterate through the PDF pages and extract text
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text += page.get_text("text")
    
    # Split the text into paragraphs
    paragraphs = [para for para in text.split('\n') if para.strip() != ""]
    return paragraphs



# Function to extract a table of contents (TOC) using GPT-3.5-turbo

# Function to extract a table of contents (TOC) using GPT-3.5-turbo

# Function to extract a table of contents (TOC) using GPT-3.5-turbo
def extract_toc_gpt(paragraphs):
    # Preprocess text for TOC
    structured_text = preprocess_text_for_toc(paragraphs)
    
    # Chunking to handle long text
    toc_chunks = split_into_chunks(structured_text, max_chars=10000)  # Adjust the chunk size as needed
    toc_list = []

    for chunk in toc_chunks:
        prompt = (
            f"Génère une table des matières à partir de ce texte. "
            f"Ce texte est un cours de droit divisé en différentes parties, sections, chapitres. "
            f"Utilise les informations de structure comme les titres, les introductions de chapitre et les phrases numérotées pour créer une table des matières complète et ordonnée. "
            f"Voici le texte :\n\n" + chunk
        )
        
        try:
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1500,  # Safe token limit
                temperature=0.3,
            )
            toc = response.choices[0].message.content.strip()
            toc_list.append(toc) 
        except Exception as e:
            st.error(f"Erreur lors de l'appel à l'API OpenAI : {e}")
            return None

    # Combine all TOC parts into one
    full_toc = "\n\n".join(toc_list)
    return full_toc

# Function to preprocess text for TOC extraction
# ... (Rest of your code) ...

# Function to split the text into chunks for summarization
def split_into_chunks(text, max_chars):
    chunks = []
    current_chunk = ""
    for para in text.split("\n\n"):
        if len(current_chunk) + len(para) > max_chars:
            chunks.append(current_chunk.strip())
            current_chunk = para
        else:
            current_chunk += " " + para
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

# Function to preprocess text for TOC extraction
def preprocess_text_for_toc(paragraphs):
    structured_text = []
    for para in paragraphs:
        # Extract titles (assuming they start with a capital letter and are followed by a colon)
        match = re.search(r"^[A-Z][^:]*:\s*", para)
        if match:
            structured_text.append(match.group(0).strip())
        else:
            # Extract sentences starting with numbers or Roman numerals
            sentences = re.split(r'(?<=[.?!])\s', para)
            for i, sentence in enumerate(sentences):
                if re.match(r"^[1-9]\.\s|^I\.\s|^II\.\s|^III\.\s", sentence):
                    structured_text.append(sentence.strip())
                    if i + 1 < len(sentences):  # Check if there's a next sentence
                        structured_text.append(sentences[i + 1].strip())  # Add the next sentence if available
                    break  # Stop processing sentences after a numbered sentence
                else:
                    # Extract potential chapter introductions (assuming they start with a capital letter and are followed by a period)
                    if re.match(r"^[A-Z].*\.", sentence):
                        structured_text.append(sentence.strip())
                        if i + 1 < len(sentences):  # Check if there's a next sentence
                            structured_text.append(sentences[i + 1].strip())  # Add the next sentence if available
                        break  # Stop processing sentences after a chapter introduction
    return "\n\n".join(structured_text)


# Function to extract a table of contents (TOC) using GPT-3.5-turbo for each chunk

# Preprocess text for PDFs to remove unwanted elements and retain paragraph structure
def preprocess_paragraphs(paragraphs):
    cleaned_paragraphs = []
    for para in paragraphs:
        para = re.sub(r"\b\w+\s\d{4}\b", "", para)  # Remove "Name Year" patterns
        para = re.sub(r"\b\d+\b", "", para)         # Remove standalone numbers
        para = re.sub(r"\s+", " ", para).strip()    # Remove excess whitespace
        if len(para.split()) > 4:  # Only keep meaningful paragraphs
            cleaned_paragraphs.append(para)
    return cleaned_paragraphs

# Function to split the text into chunks for summarization
def split_into_chunks(paragraphs, max_chars):
    chunks = []
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) > max_chars:
            chunks.append(current_chunk.strip())
            current_chunk = para
        else:
            current_chunk += " " + para
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

# Function to summarize a chunk of text, ensuring it doesn't exceed token limits
def summarize_chunk(chunk):
    prompt = (
        f"Vous êtes un assistant pédagogique chargé de condenser un cours de droit. "
        f"Faites un résumé concis et fluide en regroupant les idées dans des paragraphes cohérents. Incluez une table des matières au début cohérente avec le texte original, en vous inspirant des parties qui sont déjà numérotées dans le document."
        f"Ne pas inclure la table des matières dans le texte. Inclure la table des matières de toutes les parties au début du résumé"
        f"Conservez uniquement les points clés essentiels et faites un résumé exhaustif du texte suivant :\n\n"
        f"Texte à résumer :\n\n{chunk}"
    )

    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1800,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Erreur lors de l'appel à l'API OpenAI : {e}")
        return None

# Function to clean and format the output text
def clean_output(text):
    cleaned_text = re.sub(r'\bRésumé[: ]*', '', text)
    cleaned_text = cleaned_text.strip()
    return cleaned_text

# Function to summarize the entire document (for PDFs) with parallel processing
def summarize_entire_document_pdf(paragraphs):
    paragraphs = preprocess_paragraphs(paragraphs)
    chunks = split_into_chunks(paragraphs, max_chars=PDF_CHUNK_SIZE)
    
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        summarized_chunks = list(executor.map(summarize_chunk, chunks))
    
    summarized_chunks = [clean_output(summary) for summary in summarized_chunks if summary]
    
    return "\n\n".join(summarized_chunks)

# Function to summarize the entire document (for DOCX files)
def summarize_entire_document_docx(paragraphs):
    chunks = split_into_chunks(paragraphs, max_chars=DOCX_CHUNK_SIZE)
    
    summarized_chunks = []
    for chunk in chunks:
        summary = summarize_chunk(chunk)
        if summary:
            cleaned_summary = clean_output(summary)
            summarized_chunks.append(cleaned_summary)
    
    return "\n\n".join(summarized_chunks)

# Function to create a .docx file with the summarized text or table of contents
def create_docx_file(content, heading):
    doc = Document()
    doc.add_heading(heading, 0)
    doc.add_paragraph(content)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to create a PDF file with the summarized text or table of contents

# Function to create a PDF file using reportlab
def create_pdf_file(content):
    # Create a buffer to hold the PDF
    buffer = BytesIO()
    
    # Create a PDF canvas with A4 page size
    pdf = canvas.Canvas(buffer, pagesize=A4)
    
    # Set up font and size
    pdf.setFont("Helvetica", 10)

    # Split the content into lines
    lines = content.split("\n")

    # Start drawing text at the top of the page
    y_position = 800  # Start at the top of the page (A4 height is 842)
    
    for line in lines:
        # Draw each line, move down the y-axis
        pdf.drawString(50, y_position, line)
        y_position -= 12  # Move 12 points down per line
        
        # If we're near the bottom of the page, create a new page
        if y_position < 40:
            pdf.showPage()  # Finish the current page
            pdf.setFont("Helvetica", 10)  # Reset font for the new page
            y_position = 800  # Reset y-position for the new page
    
    # Finish the PDF
    pdf.save()
    
    # Get the value of the buffer and rewind it to the start
    buffer.seek(0)
    return buffer

# Streamlit interface
st.title("Résumé de cours ou Table des matières")

# File uploader
uploaded_file = st.file_uploader("Téléchargez un fichier de cours (.docx ou .pdf)", type=["docx", "pdf"])

# Initialize variables for paragraphs, summary, and table of contents
paragraphs = None
output = None

# Checkbox to select between generating a résumé or table of contents
generate_toc = st.checkbox("Générer une table des matières")

# Process file after it is uploaded
if uploaded_file is not None:
    # Detect file type and read accordingly
    if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        paragraphs = read_docx(uploaded_file)
        output_format = "docx"
    elif uploaded_file.type == "application/pdf":
        paragraphs = read_pdf(uploaded_file)
        output_format = "pdf"

    st.write("Aperçu des 5 premiers paragraphes :")
    st.text_area("Texte du document", "\n\n".join(paragraphs[:5]), height=300)

# Button to trigger summarization or table of contents generation
if st.button("Générer"):
    with st.spinner("Génération en cours..."):
        try:
            if generate_toc:
                toc = extract_toc_gpt(paragraphs)
                output = "\n".join(toc)
            else:
                if output_format == "pdf":
                    output = summarize_entire_document_pdf(paragraphs)
                else:
                    output = summarize_entire_document_docx(paragraphs)

            if output:
                st.success("Génération réussie !")
                st.write(output)

        except Exception as e:
            st.error(f"Une erreur s'est produite pendant la génération : {e}")

# Allow download of the generated document
if output:
    if generate_toc:
        heading = "Table des matières"
    else:
        heading = "Résumé du cours"

    if output_format == "pdf":
        pdf_file = create_pdf_file(output)
        st.download_button(
            label=f"Télécharger le {heading} en .pdf",
            data=pdf_file,
            file_name=f"{heading}.pdf",
            mime="application/pdf"
        )
    else:
        docx_file = create_docx_file(output, heading)
        st.download_button(
            label=f"Télécharger le {heading} en .docx",
            data=docx_file,
            file_name=f"{heading}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
