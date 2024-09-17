from concurrent.futures import ThreadPoolExecutor
from docx import Document
from io import BytesIO
import streamlit as st
import re
import openai
from fpdf import FPDF
import fitz  # PyMuPDF


import openai


# Set your OpenAI API key (you can also use environment variables for better security)
# --- Get OpenAI API key from user input ---
openai_api_key = st.text_input("Entrez votre clé API OpenAI :", type="password")

if openai_api_key:
    openai.api_key = openai_api_key
else:
    st.warning("Veuillez entrer votre clé API OpenAI pour continuer.")
    st.stop()  # Stop execution if no API key is provided


if openai_api_key is None:
    st.error("OpenAI API key not found. Please set it as an environment variable.")

openai.api_key = openai_api_key
# Constants# Constants
# Estimate words per page in Word
WORDS_PER_PAGE = 900  # Rough estimate, assuming standard 12pt font, normal margins
MAX_PAGES = 13
MAX_WORDS = WORDS_PER_PAGE * MAX_PAGES

# Function to read a .docx file and return a list of paragraphs
def read_docx(file):
    doc = Document(file)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    return paragraphs

# Function to read a PDF file and return a list of paragraphs
def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    
    # Iterate through the PDF pages and extract text
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
    
    # Split the text into paragraphs
    paragraphs = [para for para in text.split('\n') if para.strip() != ""]
    return paragraphs

# Function to break the text into smaller chunks (within token limit)
def split_into_chunks(paragraphs, max_chars=4000):
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        if len(current_chunk) + len(para) > max_chars:
            chunks.append(current_chunk)
            current_chunk = para
        else:
            current_chunk += "\n\n" + para
    
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks

# Function to summarize a chunk of text with word limits in mind
def summarize_chunk(chunk, target_word_count):
    prompt = (
        f"Vous êtes un assistant utile chargé de créer un résumé très bref en français. "
        f"Faites un résumé extrêmement concis de ce texte, en gardant uniquement les points clés essentiels "
        f"et en supprimant tout détail non pertinent. "
        f"Le résumé doit tenir dans une limite de mots, et doit permettre d'obtenir toutes les informations nécessaires "
        f"pour réussir un examen, en faisant environ {target_word_count} mots pour respecter une limite de pages :\n\n"
        f"Texte à résumer :\n\n{chunk}"
    )
    
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=1800,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Erreur lors de l'appel à l'API OpenAI : {e}")
        return None

# Function to clean the output text to make it look professional and structured
def clean_output(text):
    cleaned_text = re.sub(r'\bRésumé[: ]*', '', text)
    cleaned_text = cleaned_text.strip()
    return cleaned_text

# Function to apply summarization and limit total word count for 6 pages
def summarize_and_limit_length(paragraphs):
    # Split the text into chunks
    chunks = split_into_chunks(paragraphs, max_chars=3000)
    summarized_chunks = []
    total_word_count = 0

    for chunk in chunks:
        # Estimate target word count for the remaining chunks
        remaining_words = MAX_WORDS - total_word_count
        if remaining_words <= 0:
            break
        
        # Summarize the chunk with a target word count
        summary = summarize_chunk(chunk, target_word_count=remaining_words)
        if summary:
            cleaned_summary = clean_output(summary)
            chunk_word_count = len(cleaned_summary.split())

            # Add chunk summary if it doesn't exceed the maximum word count
            if total_word_count + chunk_word_count <= MAX_WORDS:
                summarized_chunks.append(cleaned_summary)
                total_word_count += chunk_word_count
            else:
                # Truncate the summary if the next chunk would exceed the limit
                truncated_summary = " ".join(cleaned_summary.split()[:remaining_words])
                summarized_chunks.append(truncated_summary)
                total_word_count += len(truncated_summary.split())
                break  # Stop once we reach the word limit

    # Combine all summarized chunks into one final summary
    return "\n\n".join(summarized_chunks)

# Function to create a .docx file with the summarized text
def create_summary_docx(summary):
    doc = Document()
    
    # Add the summary as the main content
    doc.add_heading('Résumé du cours de droit', 0)
    doc.add_paragraph(summary)
    
    # Save the document to a buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit interface
st.title("Résumé concis en français avec GPT-3.5 (Limité à 6 pages)")

# File uploader with a unique key
uploaded_file = st.file_uploader(
    "Téléchargez un fichier de cours (.docx ou .pdf)", 
    type=["docx", "pdf"],
    key="file_uploader_1"
)

if uploaded_file is not None:
    # Detect file type and read accordingly
    if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        paragraphs = read_docx(uploaded_file)
    elif uploaded_file.type == "application/pdf":
        paragraphs = read_pdf(uploaded_file)
    
    st.write("Aperçu des 5 premiers paragraphes :")
    st.text_area("Texte du document", "\n\n".join(paragraphs[:5]), height=300)

    # Initialize summary variable
    lesson_summary = None

    # Generate the summary using GPT-3.5 for the chunks, limiting the length
    if st.button("Générer un petit résumé (moins de 6 pages)"):
        with st.spinner("Génération du résumé..."):
            try:
                lesson_summary = summarize_and_limit_length(paragraphs)  # Summarize while limiting word count
                if lesson_summary:
                    st.success("Résumé généré avec succès !")
                    st.write(lesson_summary)
            except Exception as e:
                st.error(f"Une erreur s'est produite pendant la génération du résumé : {e}")
        
        # Only allow download if summarization was successful
        if lesson_summary:
            # Create the .docx file with the summary
            docx_file = create_summary_docx(lesson_summary)
            
            # Download the summary as a .docx file
            st.download_button(
                label="Télécharger le résumé en .docx",
                data=docx_file,
                file_name="petit_resume_cours.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
